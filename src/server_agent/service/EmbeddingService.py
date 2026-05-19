"""
EmbeddingService — 文档 chunking、embedding 生成与向量存储管理。

向量库：Chroma（本地持久化，路径 DATA_DIR/chroma）
Embedding 模型：通过 OpenAI-compatible API 调用（默认指向 LM Studio）

环境变量（均有默认值，与 LLM 配置独立）：
  EMBEDDING_MODEL      嵌入模型名称，默认 qwen3-embedding:8b
  EMBEDDING_API_BASE   embedding 服务地址，默认 http://localhost:11434/v1（Ollama）
  EMBEDDING_API_KEY    API key，默认 ollama
"""
from __future__ import annotations

import asyncio
import logging
import os
import pathlib
from functools import partial
from typing import List

from src.server_agent.mapper.paths import in_data

logger = logging.getLogger(__name__)

CHROMA_DIR = in_data("chroma")
CHROMA_DIR.mkdir(parents=True, exist_ok=True)

_CHUNK_SIZE = 500
_CHUNK_OVERLAP = 50
_SEPARATORS = ["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""]


# ---------------------------------------------------------------------------
# 同步工具函数（在线程池中执行）
# ---------------------------------------------------------------------------

def _get_embeddings():
    from langchain_openai import OpenAIEmbeddings
    return OpenAIEmbeddings(
        model=os.getenv("EMBEDDING_MODEL", "qwen3-embedding:8b"),
        api_key=os.getenv("EMBEDDING_API_KEY", "ollama"),
        base_url=os.getenv("EMBEDDING_API_BASE", "http://localhost:11434/v1"),
        check_embedding_ctx_length=False,
    )


def _extract_text(file_path: pathlib.Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix in (".docx", ".doc"):
        return _extract_word(file_path)
    if suffix in (".xlsx", ".xls"):
        return _extract_excel(file_path)
    return file_path.read_text(encoding="utf-8", errors="replace")


_OCR_CHARS_THRESHOLD = 30  # avg chars/page below this triggers OCR
_HEADER_MAX_RATIO = 0.4   # 某行占总页数比例超过此值视为页眉/页脚


def _remove_repeated_lines(pages_text: list[str], num_pages: int) -> str:
    """去除在大量页面中重复出现的页眉/页脚/水印行。"""
    from collections import Counter
    line_counts: Counter = Counter()
    for page in pages_text:
        seen = set()
        for line in page.splitlines():
            stripped = line.strip()
            if stripped and stripped not in seen:
                line_counts[stripped] += 1
                seen.add(stripped)

    threshold = max(2, int(num_pages * _HEADER_MAX_RATIO))
    noise_lines = {line for line, cnt in line_counts.items() if cnt >= threshold}
    if noise_lines:
        logger.info("PDF: removing %d repeated header/footer lines", len(noise_lines))

    cleaned_pages = []
    for page in pages_text:
        kept = [
            line for line in page.splitlines()
            if line.strip() not in noise_lines
        ]
        cleaned_pages.append("\n".join(kept))
    return "\n".join(cleaned_pages)


def _extract_pdf(path: pathlib.Path) -> str:
    """
    混合策略：
    1. 先用 pypdf 提取文字图层
    2. 去除高频重复的页眉/页脚/水印行
    3. 若过滤后平均每页字符数 < 阈值（扫描版 / 截图 PDF），自动降级到 OCR
    """
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    num_pages = max(len(pages_text), 1)
    total_chars = sum(len(t.strip()) for t in pages_text)
    avg_chars = total_chars / num_pages

    if avg_chars >= _OCR_CHARS_THRESHOLD:
        logger.info("PDF %s: text layer OK (%.0f chars/page)", path.name, avg_chars)
        cleaned = _remove_repeated_lines(pages_text, num_pages)
        clean_chars = len(cleaned.strip())
        if clean_chars < num_pages * _OCR_CHARS_THRESHOLD:
            logger.info(
                "PDF %s: after header removal only %d chars left, falling back to OCR",
                path.name, clean_chars,
            )
            return _extract_pdf_ocr(path)
        return cleaned

    logger.info(
        "PDF %s: sparse text (%.0f chars/page), falling back to OCR",
        path.name, avg_chars,
    )
    return _extract_pdf_ocr(path)


def _extract_pdf_ocr(path: pathlib.Path) -> str:
    """用 pypdfium2 渲染页面 + RapidOCR 识别文字（支持中文）。"""
    try:
        import pypdfium2 as pdfium
        from rapidocr_onnxruntime import RapidOCR
        import numpy as np
    except ImportError as e:
        logger.error("OCR 依赖缺失: %s  请运行 pip install rapidocr-onnxruntime pypdfium2", e)
        return ""

    ocr = RapidOCR()
    try:
        pdf = pdfium.PdfDocument(str(path))
    except Exception as e:
        logger.error("pypdfium2 打开失败 %s: %s", path.name, e)
        return ""

    all_text: list[str] = []
    for i in range(len(pdf)):
        try:
            page = pdf[i]
            bitmap = page.render(scale=2)      # 2× 分辨率提升识别率
            pil_img = bitmap.to_pil()
            result, _ = ocr(np.array(pil_img))
            if result:
                page_text = "\n".join(line[1] for line in result)
                all_text.append(page_text)
        except Exception as e:
            logger.warning("OCR 第 %d 页失败: %s", i, e)

    pdf.close()
    combined = "\n".join(all_text)
    logger.info("PDF OCR %s: 共识别 %d 字符", path.name, len(combined))
    return combined


def _extract_word(path: pathlib.Path) -> str:
    from docx import Document
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_excel(path: pathlib.Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    lines: List[str] = []
    for ws in wb.worksheets:
        lines.append(f"=== {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            line = "\t".join("" if v is None else str(v) for v in row)
            if line.strip():
                lines.append(line)
    wb.close()
    return "\n".join(lines)


def _embed_sync(
    doc_id: int,
    kb_id: int,
    file_path: pathlib.Path,
    chunk_size: int,
    chunk_overlap: int,
) -> int:
    """提取文本 → 分块 → 生成 embedding → 写入 Chroma。返回 chunk 数量。"""
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_chroma import Chroma

    text = _extract_text(file_path)
    if not text.strip():
        logger.warning("doc %d: no text extracted, skipping embedding", doc_id)
        return 0

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=_SEPARATORS,
    )
    chunks = splitter.split_text(text)
    if not chunks:
        return 0

    ids = [f"doc_{doc_id}_chunk_{i}" for i in range(len(chunks))]
    metadatas = [
        {"doc_id": doc_id, "kb_id": kb_id, "chunk_index": i}
        for i in range(len(chunks))
    ]

    vectorstore = Chroma(
        collection_name=f"kb_{kb_id}",
        embedding_function=_get_embeddings(),
        persist_directory=str(CHROMA_DIR),
    )
    vectorstore.add_texts(texts=chunks, metadatas=metadatas, ids=ids)

    logger.info("doc %d → kb %d: embedded %d chunks", doc_id, kb_id, len(chunks))
    return len(chunks)


def _delete_doc_sync(doc_id: int, kb_id: int) -> None:
    try:
        import chromadb
        client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        collection = client.get_collection(f"kb_{kb_id}")
        results = collection.get(where={"doc_id": doc_id})
        if results and results["ids"]:
            collection.delete(ids=results["ids"])
            logger.info("deleted %d chunks for doc %d", len(results["ids"]), doc_id)
    except Exception as exc:
        logger.warning("delete_doc_sync(doc=%d, kb=%d): %s", doc_id, kb_id, exc)


def _delete_kb_sync(kb_id: int) -> None:
    try:
        import chromadb
        client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        client.delete_collection(f"kb_{kb_id}")
        logger.info("deleted chroma collection kb_%d", kb_id)
    except Exception as exc:
        logger.warning("delete_kb_sync(kb=%d): %s", kb_id, exc)


def _search_sync(kb_id: int, query: str, top_k: int) -> List[dict]:
    from langchain_chroma import Chroma
    vectorstore = Chroma(
        collection_name=f"kb_{kb_id}",
        embedding_function=_get_embeddings(),
        persist_directory=str(CHROMA_DIR),
    )
    docs_and_scores = vectorstore.similarity_search_with_score(query, k=top_k)
    return [
        {
            "content": doc.page_content,
            "score": float(score),
            "doc_id": doc.metadata.get("doc_id"),
            "chunk_index": doc.metadata.get("chunk_index"),
        }
        for doc, score in docs_and_scores
    ]


# ---------------------------------------------------------------------------
# 异步门面
# ---------------------------------------------------------------------------

class EmbeddingService:
    """
    同步 Chroma 操作的异步封装，所有耗时操作均在线程池中执行，
    不会阻塞 FastAPI 事件循环。
    """

    async def embed_document(
        self,
        doc_id: int,
        kb_id: int,
        file_path: pathlib.Path,
        chunk_size: int = _CHUNK_SIZE,
        chunk_overlap: int = _CHUNK_OVERLAP,
    ) -> int:
        """返回生成的 chunk 数量。"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            partial(_embed_sync, doc_id, kb_id, file_path, chunk_size, chunk_overlap),
        )

    async def delete_document_embeddings(self, doc_id: int, kb_id: int) -> None:
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, partial(_delete_doc_sync, doc_id, kb_id))

    async def delete_kb_embeddings(self, kb_id: int) -> None:
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, partial(_delete_kb_sync, kb_id))

    async def search(
        self,
        kb_id: int,
        query: str,
        top_k: int = 5,
    ) -> List[dict]:
        """语义检索，返回 [{content, score, doc_id, chunk_index}, ...]。"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            partial(_search_sync, kb_id, query, top_k),
        )
